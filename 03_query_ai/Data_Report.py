import requests
import os
import pandas as pd
import json
from datetime import datetime
import sys
import subprocess
import re
from datetime import datetime
from openai import OpenAI

#Summarizing the content of the articles and provide trend analysis of cryptocurrency market.


#Workflow: NEWSDATA_API (pull url of any related currency articles) -> zyte API (instant extraction from any page, one parameter is all you need to extract main content of a page)
# -> OPENAI API (summarize the content of the articles and provide trend analysis of cryptocurrency market)
#-----------------------Prerequisites Function----------------------
# Try to import zyte_api, auto-install if missing. zyte_api is used to extract the content of the articles.
try:
    from zyte_api import ZyteAPI
except ImportError:
    print("=" * 60)
    print("zyte_api not found. Attempting to install...")
    print(f"Using Python: {sys.executable}")
    print("=" * 60)
    
    # Try to install using the current Python interpreter
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "zyte-api"])
        print("✅ Installation successful! Re-importing...")
        from zyte_api import ZyteAPI
        print("✅ zyte_api imported successfully!")
    except subprocess.CalledProcessError:
        print("❌ Auto-installation failed.")
        print()
        print("Please run this command manually:")
        print(f"  {sys.executable} -m pip install zyte-api")
        print()
        print("Or configure your IDE to use this Python:")
        print("  C:\\Users\\user\\AppData\\Local\\Programs\\Python\\Python314\\python.exe")
        sys.exit(1)
# Try to import docx, auto-install if missing
try:
    from docx import Document
except ImportError:
    import sys
    import subprocess
    print("=" * 60)
    print("ERROR: docx module not found!")
    print("=" * 60)
    print(f"Current Python: {sys.executable}")
    print()
    print("Attempting to install python-docx...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("✅ Installation successful! Re-importing...")
        from docx import Document
        print("✅ docx imported successfully!")
    except subprocess.CalledProcessError:
        print("❌ Auto-installation failed.")
        print()
        print("Please run this command manually:")
        print(f"  {sys.executable} -m pip install python-docx")
        sys.exit(1)
# Function to remove image descriptions from the article content
def remove_image_descriptions(text: str) -> str:
    """
    Remove image descriptions and captions from article text.
    Filters out common patterns like "This image shows...", "Photo by...", etc.
    
    Parameters:
        text: The article text to clean
    
    Returns:
        str: Cleaned text without image descriptions
    """
    
    if not text:
        return ""
    
    lines = text.split('\n')
    cleaned_lines = []
    
    # Patterns that indicate image descriptions
    image_patterns = [
        r'^This image',
        r'^This photo',
        r'^Photo by',
        r'^Image by',
        r'^Image released by',
        r'^Image via',
        r'^Photo via',
        r'^\(Photo by',
        r'^\(Image by',
        r'^\(AP Photo',
        r'^\(.*?via AP\)',  # (Something via AP)
        r'^\(.*?Photo.*?\)',  # (Something Photo Something)
        r'^Director.*?pose for photographers',  # Photo call descriptions
        r'^This image released by',
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            cleaned_lines.append('')
            continue
        
        # Skip lines that match image description patterns
        is_image_desc = False
        for pattern in image_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                is_image_desc = True
                break
        
        # Skip very short lines that are likely captions (less than 20 chars)
        # unless they look like actual content
        if not is_image_desc and len(line) < 20:
            # Keep if it looks like a sentence (ends with punctuation)
            if not re.search(r'[.!?]$', line):
                # Skip if it's all caps (likely a caption)
                if line.isupper():
                    continue
        
        if not is_image_desc:
            cleaned_lines.append(line)
    
    # Join lines and clean up multiple newlines
    cleaned_text = '\n'.join(cleaned_lines)
    # Remove excessive blank lines (more than 2 consecutive)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    return cleaned_text.strip()
#Function to format the date to YYYY-MM-DD for effective API call
def normalize_date(user_input: str) -> str:
    """
    Convert various user date inputs into 'YYYY-MM-DD' format.
    Raises ValueError if format is invalid.
    """

    # Common date formats users might type
    possible_formats = [
        "%Y-%m-%d",   # 2023-02-14
        "%Y/%m/%d",   # 2023/02/14
        "%Y.%m.%d",   # 2023.02.14
        "%d-%m-%Y",   # 14-02-2023
        "%d/%m/%Y",   # 14/02/2023
        "%m-%d-%Y",   # 02-14-2023
        "%m/%d/%Y",   # 02/14/2023
        "%Y%m%d",     # 20230214
        "%b %d %Y",   # Feb 14 2023
        "%B %d %Y",   # February 14 2023
    ]

    user_input = user_input.strip()

    for fmt in possible_formats:
        try:
            parsed_date = datetime.strptime(user_input, fmt)
            return parsed_date.strftime("%Y-%m-%d")
        except ValueError:
            continue

    raise ValueError("Invalid date format. Please enter a valid date.")
#Get full article data (includes headline, body, metadata, etc.)
def extract_article_content(url: str):
    result = client.get({
        "url": url,
        "article": True  # Extracts structured article data
    })

    # Extract only the article content (body text)
    article = result.get("article", {})
    article_content = article.get("articleBody", "")  # Just the text content
    article_content = remove_image_descriptions(article_content)
    # print("Article Content Only:")
    # print("=" * 60)
    # print(article_content)
    # print("=" * 60)
    return article_content
#Schema for the response from the Ollama model - structured output with key insights only
#Schema for the response from the Ollama model - structured output with key insights only
schema = {
    "type": "object",
    "properties": {
        "key_insights": {
            "type": "array",
            "description": "List of key insights and important points extracted from the article. Each insight should be a concise, actionable point.",
            "items": {
                "type": "string"
            },
            "minItems": 3,
            "maxItems": 7
        }
    },
    "required": ["key_insights"]
}
#----------------------Step 1: Prerequisites----------------------
# Load environment variables from .env file
def load_env_file(filepath="../.env"):
    """Load variables from .env file into environment"""
    if os.path.exists(filepath):
        with open(filepath, "r") as f:
            for line in f:
                line = line.strip() #is a string method that removes leading and Trailing spaces and Newline characters \n
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1) #split the string only at the first = sign
                    os.environ[key.strip()] = value.strip()
#load .env file
load_env_file()

#Get API key from environment variable
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")
# NYT_API_KEY = os.getenv("TEST_API_KEY")
ZYTE_API_KEY = os.getenv("ZYTE_API_KEY")
# Extract only article content using Zyte API
# According to https://python-zyte-api.readthedocs.io/en/stable/intro/basic.html
client = ZyteAPI(api_key=ZYTE_API_KEY)

#----------------------Step 2: Extract Cypto Related Articles from NEWSDATA_API----------------------
# According to https://newsdata.io/crypto-news-api
# 1. Get all the latest news on new cryptocurrency about BTC and summarize each article using Ollama model.
def get_latest_crypto_news( currency: str = "btc"):

    # Format date safely - normalize_date returns a string in "YYYY-MM-DD" format
    #formatted_date = normalize_date(date)  # Returns string like "2026-02-14"

    base_url = "https://newsdata.io/api/1/crypto"

    params = {
        "apikey": "pub_887f35a8fe784da09907cb9f61da842b",
        "coin": currency,
        #"from_date": formatted_date,  # Free plan does not support from_date parameter, so it will be ignored.API expects string, not int
        "language": "en"
    }

    response = requests.get(base_url, params=params)
    response.raise_for_status()

    data = response.json()
    data_articles = [] #list to store the articles data

    #Loop through the articles and summarize each article
    for article in data.get("results", []):
        article_title = article.get("title")  # Save title before overwriting
        article_link = article.get("link")     # Save link before overwriting
        print(f"Title: {article_title} | Date: {article.get("pubDate")}")
        print(f"Processing article: {article_title}")
        
        # Extract article content (returns a string, not a dict)
        article_content = extract_article_content(article_link)
        
        # Feed the article content to the Ollama model to summarize the content of the article
        PORT = 11434
        OLLAMA_HOST = f"http://localhost:{PORT}"
        url = f"{OLLAMA_HOST}/api/generate"
        #Build the request body as a dictionary
        body = {
            "model": "crpytoanalyst:latest",
            "prompt": article_content,  # Use article_content (string), not article (dict)
            "format":schema,
            "stream": False
        }
        #Build and send the POST request to the Ollama REST API
        response = requests.post(url, json=body)
        #Parse the response JSON
        response_data = response.json()
        output = response_data["response"]
        
        # Parse the structured output (should be JSON string)
        try:
            insights_data = json.loads(output) if isinstance(output, str) else output
            print("\n📊 Key Insights from Article:")
            print("-" * 120)
            for i, insight in enumerate(insights_data.get('key_insights', []), 1):
                print(f"{i}. {insight}")
            print("-" * 120)
            #Store the article data in a dictionary
            article_data = {
                "title": article_title,
                "date": article.get("pubDate"),
                "key_insights": insights_data.get('key_insights', []) #list of key insights
            }
            data_articles.append(article_data)
            
        except (json.JSONDecodeError, TypeError, KeyError) as e:
            # Fallback if response is not valid JSON or doesn't match schema
            print("\n📊 Article Analysis:")
            print("-" * 120)
            print(output)
            print("-" * 120)

        print(f"Processing completed for article: {article_title}")  # Use saved title
        print("=" * 120)
    return data_articles

#----------------------Step 3: Process the API data (clean, filter, or aggregate as needed)----------------------
# def normalize_nyt_person(name: str) -> str:
#     """
#     Convert NYT person facet from 'Last, First Middle' to 'First Middle Last'.
#     Keeps suffixes reasonably well (e.g., 'King Jr., Martin Luther' -> 'Martin Luther King Jr.').
#     """
#     if not name or not isinstance(name, str):
#         return name

#     name = name.strip() #remove redundant spaces

#     # If there's no comma, assume it's already in display order
#     if "," not in name:
#         return " ".join(name.split()) #split words and join them with a space

#     # Split only on the first comma
#     last, rest = name.split(",", 1)
#     last = last.strip()
#     rest = rest.strip()

#     # Simple handling for suffixes that sometimes appear with last name (rare in NYT facets but possible)
#     # e.g., "King Jr., Martin Luther" -> last = "King Jr."
#     # This already works: "Martin Luther" + "King Jr."
#     display = f"{rest} {last}".strip()

#     # Normalize whitespace
#     return " ".join(display.split()) 

# def combined_json(document):
#     """Combine document fields into a single text string for embedding"""
#     return f"Title: {document.get('title', '')} Date: {document.get('published_date', '')} Abstract: {document.get('abstract', '')}"
#----------------------Step 3: Generate a report of the article using OPENAI API----------------------
def get_data_report(articles_data):
    #Check if API key is set
    if not OPENAI_API_KEY:
        raise ValueError("OPENAI_API_KEY not found in .env file. Please set it up first.")

    # Convert articles_data to JSON string for the API
    # articles_data is a list of dictionaries, convert to JSON string
    articles_json = json.dumps(articles_data, indent=2, ensure_ascii=False)

    client = OpenAI(api_key=OPENAI_API_KEY)
    
    # Format input as messages with role (required by responses.create API)
    # The input parameter expects an array of message objects with 'role' field
    input_messages = [
        {
            "role": "user",
            "content": articles_json
        }
    ]
    
    response = client.responses.create(
        model="gpt-5",
        instructions="""You are a professional cryptocurrency analyst and writer with over 30 years of experience. You will receive multiple articles in 
        JSON format containing the title, date, and key insights of each article. Use a professional, objective, and concise writing style to produce a single consolidated investor brief with length of 300-400 words.
        The output should follow the following format:
        #Data Report
        ###Date: <date of the article is determined based on the time range mentioned in the articles_data—from the earliest time to the most recent time>
        Sources included:
        - <title of the article>
        ###Summary:
        <summary of the articles_data>""",
        input=input_messages,  # Pass as array of message objects with roles
        stream=False
    )
    return response.output_text


#----------------------Step 4: Main Function / Query the New York TimesAPI and process the data----------------------
def query_nyt_api(period: int = 1, num_articles: int = 20):
    """
    Query NYT Most Popular API.
    
    Parameters:
        period: Time period in days (must be 1, 7, or 30)
        num_articles: Number of articles to return (default: 20)
    
    Returns:
        List of article dictionaries, or empty list on error
    """
    articles_data = []
    #NYT API endpoint - period must be 1, 7, or 30 days
    nyt_base_url = "https://api.nytimes.com/svc/mostpopular/v2"
    nyt_url = f"{nyt_base_url}/viewed/{period}.json"  # Period: 1, 7, or 30 days
    params = {"api-key": NYT_API_KEY}
    response = requests.get(nyt_url, params=params)
    
    if response.status_code == 200:
        data = response.json()
        data_articles = data['results'][:num_articles]  # Limit to num_articles

        #Prepare data for saving /embedding
        for i, article in enumerate(data_articles):
            article_dict = {
                "date" : article.get('published_date', ''),
                "title" : article.get('title', ''),
                "abstract" : article.get('abstract', ''),
            }

            articles_data.append(article_dict)
        
        return articles_data
    else:
        # Return empty list on error, or you could raise an exception
        print(f"Error: API returned status code {response.status_code}")
        print(f"Response: {response.text}")
        return []

def main():
    # period must be 1, 7, or 30 (days)
    # num_articles is how many articles to return
    # articles_data = query_nyt_api(period=7, num_articles=10)
    # print(articles_data)
    #extract_article_content("https://coincu.com/news/bitcoin-trades-as-liquidity-chip-structure-steer-flows/")
    articles_data = get_latest_crypto_news(currency="btc")
    data_report = get_data_report(articles_data)

    #save the data_report to a file
    doc = Document()
    for line in data_report.split("\n"):
        if line.startswith("# "):
            # Main heading
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            # Subheading
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            # Bullet point
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            # Regular paragraph
            doc.add_paragraph(line)
    doc.save("data_report.docx")
    print("✅Data report saved to data_report.docx")

if __name__ == "__main__":
    main()